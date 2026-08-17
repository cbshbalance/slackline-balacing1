from __future__ import annotations

import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


WORK = Path(__file__).resolve().parent
SOURCE = WORK / "학생보고서_직선서술판.md"
ASSETS = WORK / "report_assets"
OUTPUT = WORK / "줄타기_작품설명서_그림포함_초안.docx"

# narrative_proposal preset + named competition-report override.
FONT_BODY = "맑은 고딕"
FONT_MATH = "Cambria Math"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "20364A"
MUTED = "666666"
LIGHT_FILL = "F4F6F9"
GOLD = "A26714"
CONTENT_DXA = 9638  # A4, 20 mm left/right margins -> 170 mm usable width.
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


FIGURES = [
    (1, "fig01_west_korea.png", "서양의 장대 줄타기와 한국의 전통 줄타기. 한국 줄타기는 장대 없이 부채만 든 채 처진 줄 위에서 역동적인 동작을 수행한다."),
    (2, "fig02_prior_study.png", "Paoletti와 Mahadevan의 선행연구: 원호 위 카트–단일 역진자 모델과 줄 처짐에 따른 제어입력의 변화."),
    (3, "fig03_rope_geometry.png", "처진 줄의 실제 기하. 발은 V자 줄의 꼭짓점에 고정되어 줄과 함께 움직이며, 발의 자취가 반경 R의 원호를 이룬다."),
    (4, "fig04_system_model.png", "처진 줄 진자 위 상·하체 2링크 모델과 결합 운동방정식. φ·α·θ는 모두 연직 기준 오른쪽을 양(+)으로 잡았다."),
    (5, "fig05_fwe_sequence.png", "FWE의 인과 사슬: 기울어진 상태에서 허리를 접어 발을 무게중심 너머로 보내고, 기다린 뒤 자세를 회복한다."),
    (6, "fig06_mode_thought_experiment.png", "입력 없이 떨어지지 않는 운동을 찾기 위한 세 후보: β=0, β=−φ, 그리고 그 사이의 일정한 비율."),
    (7, "fig07_stable_unstable_modes.png", "결합계의 두 운동: 줄과 몸이 엇갈려 유계 진동하는 안정모드와 지수적으로 커지는 불안정모드."),
    (8, "fig08_pendulum_contrast.png", "가운데로 되돌아오는 정진자와 기울수록 더 쓰러지는 역진자의 대조."),
    (9, "fig09_epsilon_geometry.png", "무게중심 등고선과 안정모드선의 교점으로 결정되는 최적 접기 비율 ε*."),
    (10, "fig10_epsilon_sim_steps.png", "초기 기울기에서 ε*를 계산하고, 접기로 상태를 안정모드선 위에 옮긴 뒤 유계 진동에 들어가는 과정."),
    (11, "fig11_raw_vs_predicted.png", "원위치 평면과 예측점 평면의 비교. 속도를 포함한 예측점 평면에서 안정모드선은 정확한 균형 판정선이 된다."),
    (12, "fig12_single_fold_bounded.png", "S1 순간 단일접기: ε*만큼 한 번 접으면 펴지 않아도 불안정 성분이 사라져 유계 진동한다."),
    (13, "fig13_full_cycle.png", "S2 순간 완전 사이클: 과접기–대기–펴기로 펴기 뒤의 상태까지 안정모드선 위에 착지시킨다."),
    (14, "fig14_finite_time.png", "S3 유한시간 접기: 접는 동안에도 불안정모드가 자라므로 순간접기보다 더 많이 접어야 한다."),
    (15, "fig15_repeated_error.png", "S5 반복 사이클: 대기시간이 길수록 같은 실행오차가 크게 증폭되며, 짧은 반복 보정이 더 강건하다."),
    (16, "fig16_simulation_evolution.png", "시뮬레이션의 진화: LQR, MuJoCo 전물리 모델, 유한시간 FWE, 증분 접기."),
    (17, "fig17_survival_comparison.png", "전물리 시뮬레이션의 생존 결과. 증분 접기는 20개 초기조건 모두 120초 상한에 도달했고, 상한 없는 후속 계산에서는 27,000초까지 생존했다."),
    (18, "fig18_robot_design_real.png", "줄타기 로봇 ‘어름이’의 CAD 설계와 제작된 실물."),
    (19, "fig19_robot_components.png", "로봇의 상체·하체·중앙 결합부 구성."),
    (20, "fig20_robot_measurements.png", "로봇의 진자 운동, 무게중심 위치 및 상·하체 물리량 실측 과정."),
]
FIGURE_BY_ID = {number: (ASSETS / filename, caption) for number, filename, caption in FIGURES}


# Insert after the paragraph containing the key. Multiple figures may share one anchor.
FIGURE_ANCHORS = [
    ("그들의 손에는 기다란 장대는커녕", [1]),
    ("제어 노력이 최소가 되는 최적의 줄 처짐", [2]),
    ("원호 위를 흔들리는 물체", [3]),
    ("이 좌표는 세 가지 좋은 성질", [4]),
    ("얼마나 접어야 하는가? 얼마나 기다렸다 펴야 하는가?", [5]),
    ("그 함수를 구할 자신이 없었다", [6]),
    ("이 직선 위의 초기조건에서 출발한 상태는 영원히 직선 위에 머문다", [7, 8]),
    ("초기 기울기 β₀가 식에서 소거된다", [9, 10]),
    ("그림 위의 길이가 된다", [11]),
    ("이후 펴지 않아도 그 직선을 따라 영원히 유계 진동한다", [12]),
    ("예측점 평면의 S2가 이 장면이다", [13]),
    ("γ는 그 사이에 불안정 모드가 자란 만큼을 접기량으로 미리 갚는 보정이다", [14]),
    ("루프이득 0.24 유계 vs 58 발산", [15]),
    ("최종 진화를 이끌었다", [16]),
    ("전물리 시뮬레이션에서 작동했음을 보인 결과다", [17]),
    ("실물 검증은 현재 진행 중이다", [18, 19, 20]),
]


DISPLAY_REPAIRS = {
    "R=L/22−D/22": "R = ½√(L² − D²)",
    "xCoM=−g\u2009φt\u2001\u2001(무게중심 보존법칙)": "d²x_CoM/dt² = −g·φ(t)    (무게중심 보존법칙)",
    "Mq=Gq+Eτ,\u2001\u2001q=φ,α,θT,\u2001E=0,\u2009−1,\u2009+1T": "M·d²q/dt² = Gq + Eτ,    q = (φ, α, θ)ᵀ,    E = [0, −1, +1]ᵀ",
    "M=MtR2Rp1Rp2Rp1JααJαθRp2JαθJθθ,\u2001\u2001G=diag−MtgR,\u2004p1g,\u2004p2g": "M = [[M_tR², Rp₁, Rp₂], [Rp₁, J_αα, J_αθ], [Rp₂, J_αθ, J_θθ]],    G = diag(−M_tgR, p₁g, p₂g)",
    "쓰러짐 각 β=p1α+p2θ/Mth — 상·하체 기울기를 질량 배분으로 가중평균한 것. “몸 전체가 실질적으로 얼마나 기울었나”": "쓰러짐 각 β = (p₁α + p₂θ)/(M_t h) — 상·하체 기울기를 질량 배분으로 가중평균한 것. “몸 전체가 실질적으로 얼마나 기울었나”",
    "접힘각 δ=θ−α — 우리가 명령하는 자세 변수": "접힘각 δ = θ − α — 우리가 명령하는 자세 변수",
    "xCoM=Rφ+hβ\u2001\u2001(h=직립 무게중심 높이)": "x_CoM = Rφ + hβ    (h = 직립 무게중심 높이)",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, *, name=FONT_BODY, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_para_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("- ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    r.append(fld_char1)
    r.append(instr_text)
    r.append(fld_char2)
    paragraph._p.append(r)
    run = paragraph.add_run(" -")
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "제72회 전국과학전람회 | 학생부·물리"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.runs[0], size=8.5, color=MUTED)
    footer = section.footer
    add_page_field(footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_BODY
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Equation Display" not in [s.name for s in doc.styles]:
        eq = doc.styles.add_style("Equation Display", WD_STYLE_TYPE.PARAGRAPH)
    else:
        eq = doc.styles["Equation Display"]
    eq.font.name = FONT_MATH
    eq._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    eq.font.size = Pt(10.5)
    eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(5)
    eq.paragraph_format.space_after = Pt(7)
    eq.paragraph_format.keep_together = True

    if "Figure Caption" not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["Figure Caption"]
    cap.font.name = FONT_BODY
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(9)
    cap.paragraph_format.keep_together = True


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("제72회 전국과학전람회 작품설명서")
    set_run_font(r, size=12, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("한국 줄타기는 어떻게\n장대 없이 균형을 잡는가?")
    set_run_font(r, size=26, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(50)
    r = p.add_run("처진 줄 진자 위 이중 역진자가 허리 접기 하나로 서는 원리의 규명과 로봇 실증")
    set_run_font(r, size=13, color=DARK_BLUE)

    metadata = [
        ("출품 부문", "학생부 · 물리"),
        ("출품 학생", "김지훈 · 송주호 · 이혜원"),
        ("지도 교사", "정도일"),
        ("작성 시점", "2026년 8월"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row, (label, value) in zip(table.rows, metadata):
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(8.5)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p0.paragraph_format.space_after = Pt(0)
        set_run_font(p0.add_run(label), size=10.5, bold=True, color=MUTED)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_run_font(p1.add_run(value), size=10.5, color=NAVY)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    p = doc.add_paragraph("차례", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = [
        "Ⅰ. 탐구 동기 및 목적",
        "Ⅱ. 탐구 과정",
        "  1. 문헌 조사 — 찾은 것은 답이 아니라 공백이었다",
        "  2. 첫 번째 열쇠 — 처진 줄은 받침이 아니라 변환기다",
        "  3. 발을 무게중심 너머로 보내 쓰러짐 방향을 바꾼다 — FWE 메커니즘",
        "  4. 완전히 정지할 필요 없이 떨어지지만 않으면 된다 — 한국 줄타기에 맞는 균형의 정의",
        "  5. 흔들림 속에서 균형의 길을 찾다 — 결합 모드와 안정모드선",
        "  6. 기운 크기가 달라도 같은 비율로 — 상태공간의 기하와 최적 접기 비율 ε*",
        "  7. 위치와 속도를 하나의 점에 담다 — 예측점 평면과 차원 축소",
        "  8. 접고, 기다리고, 펴는 박자를 닫다 — 순간동작 가정의 완전 FWE 사이클",
        "  9. 완벽한 한 번보다 불완전한 여러 번 — 유한시간 접기와 증분 접기의 시뮬레이션",
        "  10. 이론을 현실의 줄 위에 세우다 — 줄타기 로봇 ‘어름이’",
        "Ⅲ. 탐구 결론 및 의의",
        "Ⅳ. 참고 문헌",
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6 if item.startswith("  ") else 0)
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(item.strip()), size=9.8, bold=not item.startswith("  "), color=NAVY)

    doc.add_page_break()
    p = doc.add_paragraph("그림 차례", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for number, _, caption in FIGURES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(f"그림 {number}. "), size=9.2, bold=True, color=DARK_BLUE)
        set_run_font(p.add_run(caption), size=9.2, color=NAVY)
    doc.add_page_break()


def clean_inline(text: str) -> str:
    text = DISPLAY_REPAIRS.get(text, text)
    text = text.replace("**", "").replace("`", "")
    return text.strip()


def add_text_with_pending_highlight(paragraph, text: str, *, size=10.5, color=None, bold=False, italic=False) -> None:
    parts = re.split(r"(【[^】]+】)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
        if part.startswith("【") and part.endswith("】"):
            run.font.highlight_color = 7  # yellow
            run.font.color.rgb = RGBColor(156, 87, 0)


def add_figure(doc: Document, number: int) -> None:
    path, caption = FIGURE_BY_ID[number]
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as im:
        ratio = im.width / im.height
    width = Cm(16.2)
    if ratio < 1.0:
        width = Cm(12.0)
    elif ratio < 1.35:
        width = Cm(14.4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=width)
    cap = doc.add_paragraph(style="Figure Caption")
    add_text_with_pending_highlight(cap, f"그림 {number}. {caption}", size=9, color=MUTED)


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 2:
        return [2100, CONTENT_DXA - 2100]
    if cols == 3:
        return [2100, 2200, CONTENT_DXA - 4300]
    if cols == 5:
        return [2100, 1500, 1700, 1350, CONTENT_DXA - 6650]
    return [CONTENT_DXA // cols] * (cols - 1) + [CONTENT_DXA - (CONTENT_DXA // cols) * (cols - 1)]


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Cm(width / 1440 * 2.54)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    widths = table_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for cell, value in zip(row.cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or len(rows[0]) > 3 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_text_with_pending_highlight(p, clean_inline(value), size=8.6, bold=row_index == 0, color=NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def split_markdown_table_row(line: str) -> list[str]:
    # In this report |A| is an absolute-value symbol, not two table delimiters.
    protected = line.strip().strip("|").replace("|A|", "§ABS_A§")
    return [clean_inline(cell.strip().replace("§ABS_A§", "|A|")) for cell in protected.split("|")]


def add_body(doc: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    inserted: set[int] = set()
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# ") or stripped == "## 처진 줄 진자 위 이중 역진자가 허리 접기 하나로 서는 원리의 규명과 로봇 실증":
            index += 1
            continue
        if stripped.startswith("> 세 DOCX와 학생 최신본"):
            index += 1
            continue
        if stripped.startswith("그림  서양 줄타기") or stripped.startswith("그림  한국 전통 줄타기"):
            index += 1
            continue
        if stripped.startswith("[그림]") or stripped.startswith("시각자료 (") or stripped.startswith("그림 4-1."):
            index += 1
            continue

        if stripped.startswith("|"):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = split_markdown_table_row(lines[index])
                if not is_separator_row(cells):
                    rows.append(cells)
                index += 1
            if rows:
                add_markdown_table(doc, rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            text = clean_inline(heading.group(2))
            if level == 1 and text.startswith(("Ⅰ.", "Ⅱ.", "Ⅲ.", "Ⅳ.")):
                if text != "Ⅰ. 탐구 동기 및 목적":
                    doc.add_page_break()
            p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            add_text_with_pending_highlight(p, text, size={1: 16, 2: 13, 3: 12}[min(level, 3)], bold=True, color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[min(level, 3)])
            index += 1
            continue

        if stripped.startswith(">"):
            text = clean_inline(stripped[1:].strip())
            equation_like = bool(re.search(r"[=Δεφβλ]|CoM", text)) and not text.startswith("“")
            p = doc.add_paragraph(style="Equation Display" if equation_like else None)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if equation_like else WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(8)
            set_para_shading(p, LIGHT_FILL)
            add_text_with_pending_highlight(p, text, size=10.3, color=NAVY, italic=not equation_like)
            index += 1
        else:
            text = clean_inline(stripped)
            display_eq = text in DISPLAY_REPAIRS.values() or text in DISPLAY_REPAIRS.keys() or (
                len(text) < 150 and bool(re.match(r"^(R\s*=|x.?_?CoM\s*=|M q|M =)", text))
            )
            if display_eq:
                p = doc.add_paragraph(style="Equation Display")
                add_text_with_pending_highlight(p, text, size=10.5, color=NAVY)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.widow_control = True
                add_text_with_pending_highlight(p, text, size=10.5, color=None)
            index += 1

        for anchor, figure_numbers in FIGURE_ANCHORS:
            if anchor in stripped:
                for number in figure_numbers:
                    if number not in inserted:
                        add_figure(doc, number)
                        inserted.add(number)

    missing = sorted(set(FIGURE_BY_ID) - inserted)
    if missing:
        raise RuntimeError(f"Figures not inserted: {missing}")


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "한국 줄타기는 어떻게 장대 없이 균형을 잡는가?"
    props.subject = "제72회 전국과학전람회 학생부 물리 작품설명서"
    props.author = "김지훈, 송주호, 이혜원"
    props.last_modified_by = "정도일"
    props.keywords = "한국 줄타기, 처진 줄, FWE, 모드 분석, 예측점 평면, 증분 접기"


def main() -> None:
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_body(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
